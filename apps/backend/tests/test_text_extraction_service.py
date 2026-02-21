import io

import pytest

from app.services.text_extraction import TextExtractionError, extract_text_from_bytes


def _build_minimal_pdf_with_text(text: str) -> bytes:
    def pdf_obj(obj_num: int, body: bytes) -> bytes:
        return f"{obj_num} 0 obj\n".encode() + body + b"\nendobj\n"

    def escape_pdf_string(value: str) -> str:
        return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"

    stream = (
        "BT\n"
        "/F1 24 Tf\n"
        "72 200 Td\n"
        f"({escape_pdf_string(text)}) Tj\n"
        "ET\n"
    ).encode()

    obj1 = pdf_obj(1, b"<< /Type /Catalog /Pages 2 0 R >>")
    obj2 = pdf_obj(2, b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    obj3 = pdf_obj(
        3,
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
    )
    obj4 = pdf_obj(4, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    obj5_body = b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n"
    obj5_body += stream + b"endstream"
    obj5 = pdf_obj(5, obj5_body)

    objects = [obj1, obj2, obj3, obj4, obj5]
    offsets: list[int] = []
    out = bytearray(header)
    for obj in objects:
        offsets.append(len(out))
        out.extend(obj)

    xref_start = len(out)
    size = len(objects) + 1
    out.extend(f"xref\n0 {size}\n".encode())
    out.extend(b"0000000000 65535 f \n")
    for off in offsets:
        out.extend(f"{off:010d} 00000 n \n".encode())
    out.extend(b"trailer\n")
    out.extend(f"<< /Size {size} /Root 1 0 R >>\n".encode())
    out.extend(b"startxref\n")
    out.extend(f"{xref_start}\n".encode())
    out.extend(b"%%EOF\n")
    return bytes(out)


def test_extract_text_from_docx_bytes():
    pytest.importorskip("docx")
    from docx import Document

    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("Titulo")
    document.add_paragraph("1. Referencia")
    document.add_paragraph("Item bullet", style="List Bullet")
    document.save(buf)

    text = extract_text_from_bytes(
        source_type="docx",
        content=buf.getvalue(),
        max_chars=1_000_000,
    )
    assert "Titulo" in text
    assert "1. Referencia" in text
    assert "Item bullet" in text


def test_extract_text_from_docx_too_large():
    pytest.importorskip("docx")
    from docx import Document

    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("hola")
    document.save(buf)

    with pytest.raises(TextExtractionError) as exc:
        extract_text_from_bytes(source_type="docx", content=buf.getvalue(), max_chars=1)

    assert exc.value.code == "extracted_text_too_large"


def test_extract_text_from_docx_invalid_bytes():
    pytest.importorskip("docx")
    with pytest.raises(TextExtractionError) as exc:
        extract_text_from_bytes(
            source_type="docx",
            content=b"not-a-docx",
            max_chars=10_000,
        )

    assert exc.value.code == "text_extraction_failed"


def test_extract_text_from_pdf_bytes():
    pytest.importorskip("pdfminer")

    content = _build_minimal_pdf_with_text("Hello world")
    text = extract_text_from_bytes(
        source_type="pdf",
        content=content,
        max_chars=1_000_000,
    )
    assert "Hello world" in text
